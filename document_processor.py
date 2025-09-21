"""
Advanced Document Loader and Processing System for Deep Researcher Agent
Supports multiple file formats with intelligent text splitting and preprocessing
"""

import os
import json
import pandas as pd
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
import asyncio
import aiofiles
import hashlib
import mimetypes

# LangChain imports
from langchain_core.documents import Document
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    TokenTextSplitter,
)
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    CSVLoader,
    JSONLoader,
    DirectoryLoader,
)

# Additional imports for document processing
import docx
from bs4 import BeautifulSoup
import markdown
import nltk
import spacy

from config import ResearchConfig

@dataclass
class DocumentMetadata:
    """Enhanced metadata for processed documents"""
    file_path: str
    file_type: str
    file_size: int
    created_at: str
    modified_at: str
    doc_hash: str
    chunk_count: int
    language: Optional[str] = None
    author: Optional[str] = None
    title: Optional[str] = None
    keywords: Optional[List[str]] = None
    confidence_score: float = 1.0


class AdvancedDocumentProcessor:
    """Advanced document processor with multi-format support and intelligent chunking"""
    
    def __init__(self, config: ResearchConfig = None):
        self.config = config or ResearchConfig()
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.csv': self._process_csv,
            '.json': self._process_json,
            '.docx': self._process_docx,
            '.html': self._process_html,
            '.htm': self._process_html,
        }
        
        # Initialize text splitters
        self.recursive_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.CHUNK_SIZE,
            chunk_overlap=self.config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", " ", ""],
            keep_separator=True
        )
        
        self.markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[
                ("#", "Header 1"),
                ("##", "Header 2"),
                ("###", "Header 3"),
            ]
        )
        
        self.token_splitter = TokenTextSplitter(
            chunk_size=self.config.CHUNK_SIZE,
            chunk_overlap=self.config.CHUNK_OVERLAP
        )
        
        # Initialize NLP components
        self._initialize_nlp()
    
    def _initialize_nlp(self):
        """Initialize NLP components for text analysis"""
        try:
            # Download required NLTK data
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
            nltk.download('averaged_perceptron_tagger', quiet=True)
            
            # Initialize spaCy model for language detection and analysis
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except OSError:
                print("Warning: spaCy English model not found. Language detection disabled.")
                self.nlp = None
        except Exception as e:
            print(f"Warning: NLP initialization failed: {e}")
            self.nlp = None
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate MD5 hash of file for change detection"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def _extract_keywords(self, text: str, max_keywords: int = 10) -> List[str]:
        """Extract keywords from text using NLP"""
        if not self.nlp:
            return []
        
        try:
            doc = self.nlp(text[:5000])  # Limit text for processing
            keywords = []
            
            for token in doc:
                if (token.pos_ in ['NOUN', 'ADJ', 'PROPN'] and 
                    not token.is_stop and 
                    not token.is_punct and 
                    len(token.text) > 2):
                    keywords.append(token.lemma_.lower())
            
            # Get most frequent keywords
            from collections import Counter
            keyword_counts = Counter(keywords)
            return [word for word, count in keyword_counts.most_common(max_keywords)]
        except Exception:
            return []
    
    def _detect_language(self, text: str) -> Optional[str]:
        """Detect language of the text"""
        if not self.nlp:
            return None
        
        try:
            doc = self.nlp(text[:1000])  # Sample for language detection
            return doc.lang_
        except Exception:
            return None
    
    async def _process_pdf(self, file_path: str) -> List[Document]:
        """Process PDF files"""
        try:
            loader = PyPDFLoader(file_path)
            documents = await asyncio.get_event_loop().run_in_executor(None, loader.load)
            return documents
        except Exception as e:
            print(f"Error processing PDF {file_path}: {e}")
            return []
    
    async def _process_text(self, file_path: str) -> List[Document]:
        """Process plain text files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            document = Document(
                page_content=content,
                metadata={"source": file_path}
            )
            return [document]
        except Exception as e:
            print(f"Error processing text file {file_path}: {e}")
            return []
    
    async def _process_markdown(self, file_path: str) -> List[Document]:
        """Process Markdown files with header-based splitting"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            # First split by headers
            header_splits = self.markdown_splitter.split_text(content)
            
            # If no headers found, treat as regular text
            if not header_splits:
                document = Document(
                    page_content=content,
                    metadata={"source": file_path}
                )
                return [document]
            
            documents = []
            for split in header_splits:
                document = Document(
                    page_content=split.page_content,
                    metadata={**split.metadata, "source": file_path}
                )
                documents.append(document)
            
            return documents
        except Exception as e:
            print(f"Error processing Markdown file {file_path}: {e}")
            return []
    
    async def _process_csv(self, file_path: str) -> List[Document]:
        """Process CSV files"""
        try:
            df = pd.read_csv(file_path)
            documents = []
            
            for index, row in df.iterrows():
                # Convert row to text representation
                content = "\n".join([f"{col}: {row[col]}" for col in df.columns])
                
                document = Document(
                    page_content=content,
                    metadata={
                        "source": file_path,
                        "row_index": index,
                        "columns": list(df.columns)
                    }
                )
                documents.append(document)
            
            return documents
        except Exception as e:
            print(f"Error processing CSV file {file_path}: {e}")
            return []
    
    async def _process_json(self, file_path: str) -> List[Document]:
        """Process JSON files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            data = json.loads(content)
            
            # Handle different JSON structures
            if isinstance(data, list):
                documents = []
                for i, item in enumerate(data):
                    document = Document(
                        page_content=json.dumps(item, indent=2),
                        metadata={
                            "source": file_path,
                            "item_index": i
                        }
                    )
                    documents.append(document)
                return documents
            else:
                document = Document(
                    page_content=json.dumps(data, indent=2),
                    metadata={"source": file_path}
                )
                return [document]
        except Exception as e:
            print(f"Error processing JSON file {file_path}: {e}")
            return []
    
    async def _process_docx(self, file_path: str) -> List[Document]:
        """Process DOCX files"""
        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n".join(paragraphs)
            
            document = Document(
                page_content=content,
                metadata={"source": file_path}
            )
            return [document]
        except Exception as e:
            print(f"Error processing DOCX file {file_path}: {e}")
            return []
    
    async def _process_html(self, file_path: str) -> List[Document]:
        """Process HTML files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Extract text content
            text_content = soup.get_text()
            
            # Extract metadata
            title = soup.find('title')
            title_text = title.text if title else None
            
            meta_description = soup.find('meta', attrs={'name': 'description'})
            description = meta_description.get('content') if meta_description else None
            
            document = Document(
                page_content=text_content,
                metadata={
                    "source": file_path,
                    "title": title_text,
                    "description": description
                }
            )
            return [document]
        except Exception as e:
            print(f"Error processing HTML file {file_path}: {e}")
            return []
    
    def _create_document_metadata(self, file_path: str, documents: List[Document]) -> DocumentMetadata:
        """Create comprehensive metadata for a document"""
        file_stat = os.stat(file_path)
        file_hash = self._calculate_file_hash(file_path)
        
        # Extract sample text for analysis
        sample_text = ""
        for doc in documents[:3]:  # Sample from first few chunks
            sample_text += doc.page_content[:1000] + " "
        
        keywords = self._extract_keywords(sample_text)
        language = self._detect_language(sample_text)
        
        return DocumentMetadata(
            file_path=file_path,
            file_type=Path(file_path).suffix.lower(),
            file_size=file_stat.st_size,
            created_at=str(file_stat.st_ctime),
            modified_at=str(file_stat.st_mtime),
            doc_hash=file_hash,
            chunk_count=len(documents),
            language=language,
            keywords=keywords
        )
    
    async def process_file(self, file_path: str) -> tuple[List[Document], DocumentMetadata]:
        """Process a single file and return documents with metadata"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.stat().st_size > self.config.MAX_DOCUMENT_SIZE:
            raise ValueError(f"File too large: {file_path} (max: {self.config.MAX_DOCUMENT_SIZE} bytes)")
        
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {extension}")
        
        # Process the file
        processor_func = self.supported_extensions[extension]
        documents = await processor_func(str(file_path))
        
        if not documents:
            raise ValueError(f"No content extracted from file: {file_path}")
        
        # Split large documents into chunks
        final_documents = []
        for doc in documents:
            if len(doc.page_content) > self.config.CHUNK_SIZE:
                chunks = self.recursive_splitter.split_documents([doc])
                final_documents.extend(chunks)
            else:
                final_documents.append(doc)
        
        # Add unique IDs to documents
        for i, doc in enumerate(final_documents):
            doc.metadata["chunk_id"] = f"{file_path.stem}_{i}"
            doc.metadata["file_path"] = str(file_path)
        
        # Create metadata
        metadata = self._create_document_metadata(str(file_path), final_documents)
        
        return final_documents, metadata
    
    async def process_directory(self, directory_path: str, recursive: bool = True) -> tuple[List[Document], List[DocumentMetadata]]:
        """Process all supported files in a directory"""
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        pattern = "**/*" if recursive else "*"
        files = [f for f in directory_path.glob(pattern) 
                if f.is_file() and f.suffix.lower() in self.supported_extensions]
        
        all_documents = []
        all_metadata = []
        
        for file_path in files:
            try:
                documents, metadata = await self.process_file(str(file_path))
                all_documents.extend(documents)
                all_metadata.append(metadata)
                print(f"Processed: {file_path} ({len(documents)} chunks)")
            except Exception as e:
                print(f"Failed to process {file_path}: {e}")
        
        return all_documents, all_metadata
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.supported_extensions.keys())


# Usage example and testing
if __name__ == "__main__":
    async def main():
        processor = AdvancedDocumentProcessor()
        
        # Test with the existing CSV file
        try:
            documents, metadata = await processor.process_file("realistic_restaurant_reviews.csv")
            print(f"Processed {len(documents)} documents")
            print(f"Metadata: {metadata}")
            
            # Print first document as example
            if documents:
                print(f"\nFirst document:\n{documents[0].page_content[:500]}...")
                print(f"Metadata: {documents[0].metadata}")
        
        except Exception as e:
            print(f"Error: {e}")
    
    # Run the async function
    asyncio.run(main())
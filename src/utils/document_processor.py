"""
Document processing utilities for file parsing, chunking, and metadata extraction
"""
import io
import mimetypes
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, BinaryIO
from datetime import datetime

import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import fitz  # PyMuPDF for advanced PDF processing

from langchain_core.documents import Document
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    CharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    TokenTextSplitter
)

from ..utils.logger import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """
    Comprehensive document processing for various file formats.
    Handles file parsing, text extraction, chunking, and metadata preservation.
    """
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        max_file_size_mb: int = 50,
        supported_formats: List[str] = None
    ):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Size of text chunks for embedding
            chunk_overlap: Overlap between consecutive chunks
            max_file_size_mb: Maximum file size in MB
            supported_formats: List of supported file extensions
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_file_size_mb = max_file_size_mb
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
        
        self.supported_formats = supported_formats or ['.pdf', '.txt', '.docx', '.md', '.csv', '.json']
        
        # Initialize text splitters
        self.recursive_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
        
        self.markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[
                ("#", "Header 1"),
                ("##", "Header 2"),
                ("###", "Header 3"),
            ]
        )
        
        logger.info(f"Initialized DocumentProcessor with chunk_size={chunk_size}, chunk_overlap={chunk_overlap}")
    
    def validate_file(self, file_path: Union[str, Path], file_content: Optional[bytes] = None) -> bool:
        """
        Validate if a file can be processed.
        
        Args:
            file_path: Path to the file
            file_content: Optional file content bytes
            
        Returns:
            True if file is valid, False otherwise
        """
        try:
            file_path = Path(file_path)
            
            # Check file extension
            if file_path.suffix.lower() not in self.supported_formats:
                logger.warning(f"Unsupported file format: {file_path.suffix}")
                return False
            
            # Check file size
            if file_content:
                file_size = len(file_content)
            else:
                if not file_path.exists():
                    logger.error(f"File does not exist: {file_path}")
                    return False
                file_size = file_path.stat().st_size
            
            if file_size > self.max_file_size_bytes:
                logger.warning(f"File too large: {file_size / (1024*1024):.1f}MB > {self.max_file_size_mb}MB")
                return False
            
            logger.debug(f"File validation passed: {file_path}")
            return True
            
        except Exception as e:
            logger.error(f"File validation failed: {e}")
            return False
    
    def extract_text_from_pdf(self, file_path: Union[str, Path, BinaryIO]) -> Dict[str, Any]:
        """
        Extract text from PDF file using both PyPDF2 and PyMuPDF for better coverage.
        
        Args:
            file_path: Path to PDF file or file-like object
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            text_content = ""
            metadata = {
                "file_type": "pdf",
                "pages": 0,
                "extraction_method": "hybrid"
            }
            
            # Try PyMuPDF first (better text extraction)
            try:
                if hasattr(file_path, 'read'):
                    # Handle file-like object
                    pdf_bytes = file_path.read()
                    if hasattr(file_path, 'seek'):
                        file_path.seek(0)  # Reset for potential second use
                    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                else:
                    doc = fitz.open(file_path)
                
                for page_num in range(doc.page_count):
                    page = doc.page(page_num)
                    page_text = page.get_text()
                    if page_text.strip():
                        text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                
                metadata["pages"] = doc.page_count
                metadata["extraction_method"] = "pymupdf"
                doc.close()
                
                if text_content.strip():
                    logger.debug(f"Successfully extracted text using PyMuPDF: {len(text_content)} characters")
                    metadata["text_length"] = len(text_content)
                    return {"content": text_content, "metadata": metadata}
            
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                if hasattr(file_path, 'read'):
                    # Handle file-like object
                    reader = PdfReader(file_path)
                else:
                    reader = PdfReader(file_path)
                
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                
                metadata["pages"] = len(reader.pages)
                metadata["extraction_method"] = "pypdf2"
                
                if text_content.strip():
                    logger.debug(f"Successfully extracted text using PyPDF2: {len(text_content)} characters")
                    metadata["text_length"] = len(text_content)
                    return {"content": text_content, "metadata": metadata}
            
            except Exception as e:
                logger.error(f"PyPDF2 extraction also failed: {e}")
            
            # If both methods fail
            raise Exception("Both PyMuPDF and PyPDF2 failed to extract text")
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            return {
                "content": "",
                "metadata": {
                    "file_type": "pdf",
                    "pages": 0,
                    "extraction_method": "failed",
                    "error": str(e)
                }
            }
    
    def extract_text_from_docx(self, file_path: Union[str, Path, BinaryIO]) -> Dict[str, Any]:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file or file-like object
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            doc = DocxDocument(file_path)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            metadata = {
                "file_type": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "text_length": len(text_content)
            }
            
            logger.debug(f"Successfully extracted DOCX text: {len(text_content)} characters")
            return {"content": text_content, "metadata": metadata}
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            return {
                "content": "",
                "metadata": {
                    "file_type": "docx",
                    "error": str(e)
                }
            }
    
    def extract_text_from_plain_text(self, file_path: Union[str, Path, BinaryIO], encoding: str = 'utf-8') -> Dict[str, Any]:
        """
        Extract text from plain text files (TXT, MD).
        
        Args:
            file_path: Path to text file or file-like object
            encoding: Text encoding
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            if hasattr(file_path, 'read'):
                # Handle file-like object
                if hasattr(file_path, 'mode') and 'b' in getattr(file_path, 'mode', ''):
                    content = file_path.read().decode(encoding)
                else:
                    content = file_path.read()
                    if isinstance(content, bytes):
                        content = content.decode(encoding)
                file_extension = getattr(file_path, 'name', '').split('.')[-1] if hasattr(file_path, 'name') else 'txt'
            else:
                file_path = Path(file_path)
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                file_extension = file_path.suffix.lower()[1:]  # Remove the dot
            
            metadata = {
                "file_type": file_extension,
                "text_length": len(content),
                "encoding": encoding
            }
            
            logger.debug(f"Successfully extracted plain text: {len(content)} characters")
            return {"content": content, "metadata": metadata}
            
        except UnicodeDecodeError as e:
            # Try alternative encodings
            for alt_encoding in ['latin-1', 'cp1252', 'utf-16']:
                try:
                    if hasattr(file_path, 'read'):
                        file_path.seek(0)
                        content = file_path.read()
                        if isinstance(content, bytes):
                            content = content.decode(alt_encoding)
                    else:
                        with open(file_path, 'r', encoding=alt_encoding) as f:
                            content = f.read()
                    
                    metadata = {
                        "file_type": file_extension if 'file_extension' in locals() else 'txt',
                        "text_length": len(content),
                        "encoding": alt_encoding
                    }
                    
                    logger.warning(f"Successfully read with alternative encoding: {alt_encoding}")
                    return {"content": content, "metadata": metadata}
                    
                except Exception:
                    continue
            
            logger.error(f"Failed to decode text file with any encoding: {e}")
            return {
                "content": "",
                "metadata": {
                    "file_type": file_extension if 'file_extension' in locals() else 'txt',
                    "error": f"Encoding error: {e}"
                }
            }
        
        except Exception as e:
            logger.error(f"Plain text extraction failed: {e}")
            return {
                "content": "",
                "metadata": {
                    "file_type": file_extension if 'file_extension' in locals() else 'txt',
                    "error": str(e)
                }
            }
    
    def extract_text_from_csv(self, file_path: Union[str, Path, BinaryIO]) -> Dict[str, Any]:
        """
        Extract text from CSV file.
        
        Args:
            file_path: Path to CSV file or file-like object
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to structured text
            text_content = f"CSV Data Summary:\n"
            text_content += f"Columns: {', '.join(df.columns.tolist())}\n"
            text_content += f"Shape: {df.shape[0]} rows × {df.shape[1]} columns\n\n"
            
            # Add column descriptions
            for col in df.columns:
                text_content += f"Column '{col}':\n"
                text_content += f"  Data type: {df[col].dtype}\n"
                text_content += f"  Non-null count: {df[col].count()}\n"
                
                # Add sample values
                unique_values = df[col].dropna().unique()
                if len(unique_values) <= 10:
                    text_content += f"  Unique values: {', '.join(map(str, unique_values))}\n"
                else:
                    text_content += f"  Sample values: {', '.join(map(str, unique_values[:5]))}\n"
                text_content += "\n"
            
            # Add first few rows as text
            text_content += "Sample Data:\n"
            text_content += df.head(10).to_string(index=False)
            
            metadata = {
                "file_type": "csv",
                "rows": df.shape[0],
                "columns": df.shape[1],
                "column_names": df.columns.tolist(),
                "text_length": len(text_content)
            }
            
            logger.debug(f"Successfully extracted CSV text: {len(text_content)} characters")
            return {"content": text_content, "metadata": metadata}
            
        except Exception as e:
            logger.error(f"CSV text extraction failed: {e}")
            return {
                "content": "",
                "metadata": {
                    "file_type": "csv",
                    "error": str(e)
                }
            }
    
    def process_file(
        self,
        file_path: Union[str, Path, BinaryIO],
        filename: Optional[str] = None,
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Process a file and return chunked documents.
        
        Args:
            file_path: Path to file or file-like object
            filename: Optional filename for file-like objects
            additional_metadata: Additional metadata to include
            
        Returns:
            List of LangChain Document objects
        """
        try:
            # Determine file type
            if hasattr(file_path, 'read'):
                # Handle file-like object
                if filename:
                    file_extension = Path(filename).suffix.lower()
                else:
                    # Try to detect from content
                    file_extension = '.txt'  # Default fallback
            else:
                file_path = Path(file_path)
                filename = file_path.name
                file_extension = file_path.suffix.lower()
            
            # Validate file
            if not self.validate_file(filename or 'unknown', 
                                     file_path.read() if hasattr(file_path, 'read') else None):
                logger.error(f"File validation failed for: {filename}")
                return []
            
            # Reset file pointer if it's a file-like object
            if hasattr(file_path, 'seek'):
                file_path.seek(0)
            
            # Extract text based on file type
            if file_extension == '.pdf':
                extraction_result = self.extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                extraction_result = self.extract_text_from_docx(file_path)
            elif file_extension in ['.txt', '.md']:
                extraction_result = self.extract_text_from_plain_text(file_path)
            elif file_extension == '.csv':
                extraction_result = self.extract_text_from_csv(file_path)
            else:
                logger.error(f"Unsupported file type: {file_extension}")
                return []
            
            content = extraction_result["content"]
            file_metadata = extraction_result["metadata"]
            
            if not content.strip():
                logger.warning(f"No content extracted from file: {filename}")
                return []
            
            # Prepare base metadata
            base_metadata = {
                "filename": filename or 'unknown',
                "file_extension": file_extension,
                "processed_at": datetime.now().isoformat(),
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap
            }
            
            # Merge all metadata
            base_metadata.update(file_metadata)
            if additional_metadata:
                base_metadata.update(additional_metadata)
            
            # Split content into chunks
            if file_extension == '.md':
                # Use markdown-aware splitting
                md_docs = self.markdown_splitter.split_text(content)
                chunks = []
                for doc in md_docs:
                    # Further split large markdown sections
                    sub_chunks = self.recursive_splitter.split_text(doc.page_content)
                    for chunk in sub_chunks:
                        chunks.append(chunk)
            else:
                # Use recursive character splitting
                chunks = self.recursive_splitter.split_text(content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = base_metadata.copy()
                chunk_metadata.update({
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "chunk_id": f"{filename}_{i}" if filename else f"unknown_{i}"
                })
                
                document = Document(
                    page_content=chunk,
                    metadata=chunk_metadata
                )
                documents.append(document)
            
            logger.info(f"Successfully processed file '{filename}': {len(documents)} chunks created")
            return documents
            
        except Exception as e:
            logger.error(f"File processing failed for '{filename}': {e}")
            return []
    
    def process_multiple_files(
        self,
        file_paths: List[Union[str, Path]],
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Process multiple files and return all chunked documents.
        
        Args:
            file_paths: List of file paths
            additional_metadata: Additional metadata to include for all files
            
        Returns:
            List of all Document objects from all files
        """
        all_documents = []
        
        logger.info(f"Processing {len(file_paths)} files")
        
        for file_path in file_paths:
            try:
                documents = self.process_file(file_path, additional_metadata=additional_metadata)
                all_documents.extend(documents)
            except Exception as e:
                logger.error(f"Failed to process file {file_path}: {e}")
                continue
        
        logger.info(f"Successfully processed {len(file_paths)} files, created {len(all_documents)} total chunks")
        return all_documents
    
    def get_text_statistics(self, text: str) -> Dict[str, Any]:
        """
        Get statistics about a text string.
        
        Args:
            text: Text to analyze
            
        Returns:
            Dictionary with text statistics
        """
        lines = text.split('\n')
        words = text.split()
        
        stats = {
            "character_count": len(text),
            "word_count": len(words),
            "line_count": len(lines),
            "average_word_length": sum(len(word) for word in words) / len(words) if words else 0,
            "average_line_length": sum(len(line) for line in lines) / len(lines) if lines else 0,
            "estimated_chunks": (len(text) // self.chunk_size) + 1
        }
        
        return stats